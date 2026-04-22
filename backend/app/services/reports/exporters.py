"""Export career reports from Markdown to PDF and DOCX with proper formatting."""

from __future__ import annotations

from html import escape
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ---------------------------------------------------------------------------
# Shared Markdown → structured-blocks parser
# ---------------------------------------------------------------------------

def _parse_markdown_blocks(md: str) -> list[dict]:
    """Split markdown into structured blocks: headings, bullets, paragraphs, rules."""
    blocks: list[dict] = []
    for line in md.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            blocks.append({"type": "hr"})
            continue

        # Heading (## Title or ### Title)
        h_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            blocks.append({"type": "heading", "level": level, "text": text})
            continue

        # Unordered list item
        if re.match(r"^[-*+]\s+", stripped):
            text = re.sub(r"^[-*+]\s+", "", stripped)
            blocks.append({"type": "bullet", "text": text})
            continue

        # Ordered list item
        o_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if o_match:
            blocks.append({"type": "bullet", "text": o_match.group(1)})
            continue

        # Block quote
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            blocks.append({"type": "quote", "text": text})
            continue

        # Regular paragraph
        blocks.append({"type": "paragraph", "text": stripped})
    return blocks


def _strip_inline_md(text: str) -> str:
    """Remove inline markdown markers (bold, italic, code) for plain text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _inline_markdown_to_html(text: str) -> str:
    """Convert inline markdown to safe HTML."""
    parts = re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
    html_parts: list[str] = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            html_parts.append(f"<strong>{escape(part[2:-2])}</strong>")
        elif part.startswith("*") and part.endswith("*"):
            html_parts.append(f"<em>{escape(part[1:-1])}</em>")
        elif part.startswith("`") and part.endswith("`"):
            html_parts.append(f"<code>{escape(part[1:-1])}</code>")
        else:
            html_parts.append(escape(part))
    return "".join(html_parts)


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

_font_cache: str | None = None


def _ensure_font() -> str:
    global _font_cache
    if _font_cache:
        return _font_cache

    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in candidates:
        font_path = Path(path)
        if font_path.exists():
            try:
                pdfmetrics.registerFont(TTFont("CareerPilotFont", str(font_path)))
                _font_cache = "CareerPilotFont"
                return _font_cache
            except Exception:
                continue
    _font_cache = "Helvetica"
    return _font_cache


def _md_to_reportlab_inline(text: str) -> str:
    """Convert inline markdown bold/italic to reportlab XML tags."""
    # Bold
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # Italic
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    # Code
    text = re.sub(r"`(.+?)`", r"<font name='Courier'>\1</font>", text)
    return text


def export_markdown_to_pdf(markdown_content: str, output_path: Path) -> None:
    font_name = _ensure_font()

    styles = {
        "title": ParagraphStyle(
            "Title", fontName=font_name, fontSize=18, leading=26,
            spaceAfter=8 * mm, alignment=1,
        ),
        "h1": ParagraphStyle(
            "H1", fontName=font_name, fontSize=16, leading=22,
            spaceBefore=6 * mm, spaceAfter=4 * mm,
        ),
        "h2": ParagraphStyle(
            "H2", fontName=font_name, fontSize=14, leading=20,
            spaceBefore=5 * mm, spaceAfter=3 * mm,
        ),
        "h3": ParagraphStyle(
            "H3", fontName=font_name, fontSize=12, leading=18,
            spaceBefore=4 * mm, spaceAfter=2 * mm,
        ),
        "body": ParagraphStyle(
            "Body", fontName=font_name, fontSize=10, leading=16,
            spaceAfter=2 * mm,
        ),
        "bullet": ParagraphStyle(
            "Bullet", fontName=font_name, fontSize=10, leading=16,
            leftIndent=12 * mm, spaceAfter=1 * mm,
        ),
        "quote": ParagraphStyle(
            "Quote", fontName=font_name, fontSize=10, leading=16,
            leftIndent=8 * mm, textColor="grey",
        ),
    }

    blocks = _parse_markdown_blocks(markdown_content)
    story: list = []

    for block in blocks:
        btype = block["type"]

        if btype == "hr":
            story.append(Spacer(1, 3 * mm))
            continue

        text = block["text"]
        if btype == "heading":
            level = block["level"]
            if level == 1:
                style = styles["title"] if len(story) == 0 else styles["h1"]
            elif level == 2:
                style = styles["h2"]
            else:
                style = styles["h3"]
            story.append(Paragraph(_md_to_reportlab_inline(text), style))
            continue

        if btype == "bullet":
            safe = _md_to_reportlab_inline(f"- {text}")
            story.append(Paragraph(safe, styles["bullet"]))
            continue

        if btype == "quote":
            story.append(Paragraph(_md_to_reportlab_inline(text), styles["quote"]))
            continue

        # paragraph
        story.append(Paragraph(_md_to_reportlab_inline(text), styles["body"]))

    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )
    doc.build(story)


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def export_markdown_to_docx(markdown_content: str, output_path: Path) -> None:
    doc = Document()

    # Set default font size
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(10.5)
    font.name = "Microsoft YaHei"

    blocks = _parse_markdown_blocks(markdown_content)

    for block in blocks:
        btype = block["type"]

        if btype == "hr":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("—")
            continue

        text = block["text"]
        if btype == "heading":
            level = block["level"]
            heading_level = min(level, 4)  # docx supports 1-4 practically
            p = doc.add_heading(level=heading_level)
            _add_inline_runs(p, text)
            continue

        if btype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, text)
            continue

        if btype == "quote":
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run(text)
            run.italic = True
            continue

        # paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, text)

    doc.save(str(output_path))


def export_markdown_to_html(markdown_content: str, output_path: Path) -> None:
    """Render markdown content as a browser-friendly HTML preview."""
    blocks = _parse_markdown_blocks(markdown_content)
    title = "CareerPilot 报告预览"
    body_parts: list[str] = []
    in_list = False

    for block in blocks:
        btype = block["type"]
        text = block.get("text", "")

        if btype != "bullet" and in_list:
            body_parts.append("</ul>")
            in_list = False

        if btype == "hr":
            body_parts.append("<hr />")
            continue

        if btype == "heading":
            level = min(block["level"], 4)
            if level == 1 and title == "CareerPilot 报告预览":
                title = _strip_inline_md(text) or title
            body_parts.append(f"<h{level}>{_inline_markdown_to_html(text)}</h{level}>")
            continue

        if btype == "bullet":
            if not in_list:
                body_parts.append("<ul>")
                in_list = True
            body_parts.append(f"<li>{_inline_markdown_to_html(text)}</li>")
            continue

        if btype == "quote":
            body_parts.append(f"<blockquote>{_inline_markdown_to_html(text)}</blockquote>")
            continue

        body_parts.append(f"<p>{_inline_markdown_to_html(text)}</p>")

    if in_list:
        body_parts.append("</ul>")

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{escape(title)}</title>
  <style>
    :root {{
      color-scheme: light;
      --bg: #f5f8fc;
      --panel: #ffffff;
      --border: rgba(15, 23, 42, 0.08);
      --text: #0f172a;
      --muted: #475569;
      --accent: #0f74da;
      --quote-bg: #f8fbff;
      --quote-border: rgba(15, 116, 218, 0.24);
      --code-bg: #eff6ff;
    }}
    * {{
      box-sizing: border-box;
    }}
    body {{
      margin: 0;
      padding: 32px 16px;
      background: linear-gradient(180deg, #f7fbff 0%, var(--bg) 100%);
      color: var(--text);
      font: 16px/1.8 "Microsoft YaHei", "PingFang SC", "Noto Sans SC", sans-serif;
    }}
    main {{
      max-width: 920px;
      margin: 0 auto;
      padding: 40px 48px;
      background: var(--panel);
      border: 1px solid var(--border);
      border-radius: 20px;
      box-shadow: 0 18px 40px rgba(15, 23, 42, 0.08);
    }}
    h1, h2, h3, h4 {{
      line-height: 1.4;
      margin: 1.25em 0 0.6em;
    }}
    h1 {{
      margin-top: 0;
      font-size: 2rem;
      color: var(--accent);
    }}
    h2 {{
      font-size: 1.45rem;
    }}
    h3 {{
      font-size: 1.15rem;
    }}
    p, ul, blockquote {{
      margin: 0 0 1em;
    }}
    ul {{
      padding-left: 1.5em;
    }}
    li + li {{
      margin-top: 0.35em;
    }}
    blockquote {{
      padding: 14px 16px;
      border-left: 4px solid var(--quote-border);
      background: var(--quote-bg);
      color: var(--muted);
      border-radius: 10px;
    }}
    code {{
      padding: 0.12em 0.4em;
      border-radius: 6px;
      background: var(--code-bg);
      font-family: "Consolas", "Courier New", monospace;
      font-size: 0.92em;
    }}
    hr {{
      border: 0;
      border-top: 1px solid rgba(148, 163, 184, 0.35);
      margin: 1.5em 0;
    }}
    @media (max-width: 720px) {{
      body {{
        padding: 16px 10px;
      }}
      main {{
        padding: 24px 18px;
        border-radius: 14px;
      }}
    }}
  </style>
</head>
<body>
  <main>
    {''.join(body_parts)}
  </main>
</body>
</html>
"""
    output_path.write_text(html, encoding="utf-8")


def _add_inline_runs(paragraph, text: str) -> None:
    """Parse bold/italic/code markers in *text* and add corresponding runs."""
    # Split on **bold**, *italic*, `code` patterns
    parts = re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)
