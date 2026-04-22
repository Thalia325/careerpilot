from __future__ import annotations

import asyncio
import json
import logging
import re
import base64
from io import BytesIO
from abc import ABC, abstractmethod
from typing import Any, Optional

import httpx

logger = logging.getLogger(__name__)


# --- OCR Error Classification ---

class OCRError(Exception):
    """Base exception for OCR errors with structured metadata."""

    def __init__(self, error_code: str, message: str, retryable: bool = False) -> None:
        self.error_code = error_code
        self.message = message
        self.retryable = retryable
        super().__init__(message)

    def to_dict(self) -> dict[str, Any]:
        return {
            "error_code": self.error_code,
            "message": self.message,
            "retryable": self.retryable,
        }


class OCRParseError(OCRError):
    """File cannot be parsed (empty text, corrupt format, unsupported type)."""

    def __init__(self, message: str = "文件无法解析，请上传支持格式的文件") -> None:
        super().__init__(error_code="parse_error", message=message, retryable=False)


class OCRServiceError(OCRError):
    """Third-party OCR service returned an error (5xx, etc.)."""

    def __init__(self, message: str = "OCR 服务异常，请稍后重试") -> None:
        super().__init__(error_code="service_error", message=message, retryable=True)


class OCRNetworkError(OCRError):
    """Network or timeout error contacting OCR service."""

    def __init__(self, message: str = "网络连接超时，请检查网络后重试") -> None:
        super().__init__(error_code="network_error", message=message, retryable=True)


class BaseOCRProvider(ABC):
    @abstractmethod
    async def parse_document(
        self,
        file_name: str,
        content_bytes: bytes,
        document_type: str = "resume",
        raw_text: Optional[str] = None,
    ) -> dict[str, Any]:
        raise NotImplementedError


def _extract_keywords(text: str, candidates: list[str]) -> list[str]:
    normalized_text = _normalize_for_match(text)
    return [
        item for item in candidates
        if _normalize_for_match(item) in normalized_text
    ]


def _normalize_ocr_text(text: str) -> str:
    # Some PDF text layers are extracted as "P y t h o n" or "数 据 分 析".
    # Collapse spaces inside continuous English/digit/Chinese tokens before parsing.
    text = re.sub(r"(?<=[A-Za-z0-9])[ \t]+(?=[A-Za-z0-9])", "", text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])[ \t]+(?=[\u4e00-\u9fff])", "", text)
    return text


def _normalize_for_match(text: str) -> str:
    return re.sub(r"\s+", "", _normalize_ocr_text(text).lower())


def ocr_result_needs_refresh(ocr: dict[str, Any] | None) -> bool:
    if not ocr:
        return True
    raw_text = str(ocr.get("raw_text") or "").lstrip()
    structured = ocr.get("structured_json") or {}
    if not raw_text:
        return not isinstance(structured, dict) or not any(
            structured.get(key)
            for key in (
                "name",
                "school",
                "major",
                "grade",
                "graduation_year",
                "target_job",
                "skills",
                "projects",
                "internships",
                "certificates",
                "competitions",
                "self_evaluation",
                "gpa",
            )
        )
    head = raw_text[:2000]
    return (
        head.startswith("%PDF")
        or head.startswith("PK")
        or "\x00" in head
        or head.count("\ufffd") > 10
        or (
            not structured.get("skills")
            and bool(re.search(r"(?i)p\s*y\s*t\s*h\s*o\s*n|j\s*a\s*v\s*a|s\s*q\s*l|e\s*x\s*c\s*e\s*l", head))
        )
    )


def is_reusable_ocr_result(ocr: dict[str, Any] | None) -> bool:
    if not isinstance(ocr, dict):
        return False
    raw_text = ocr.get("raw_text")
    structured = ocr.get("structured_json")
    if not isinstance(raw_text, str) or not raw_text.strip():
        return False
    if not isinstance(structured, dict):
        return False
    return not ocr_result_needs_refresh(ocr)


def _extract_gpa(text: str) -> Optional[float]:
    match = re.search(r"(GPA|绩点)[:： ]*([0-9]\.?[0-9]*)", text, flags=re.IGNORECASE)
    return float(match.group(2)) if match else None


def _extract_major(text: str) -> str:
    explicit = re.search(r"(专业|Major)\s*[:：]\s*([^\n|｜，,；;]+)", text, flags=re.IGNORECASE)
    if explicit:
        return explicit.group(2).strip()

    known_majors = [
        "计算机科学与技术",
        "数据科学与大数据技术",
        "软件工程",
        "人工智能",
        "网络工程",
        "信息管理与信息系统",
        "数字媒体技术",
    ]
    for major in known_majors:
        if re.search(rf"{re.escape(major)}\s*专业", text):
            return major
    for major in known_majors:
        if major in text:
            return major

    degree_line = re.search(r"([^\n|｜]{2,30})\s*[|｜]\s*(本科|硕士|博士|专科)", text)
    if degree_line:
        candidate = degree_line.group(1).strip()
        if not any(token in candidate for token in ["GPA", "电话", "邮箱", "城市", "薪资"]):
            return candidate
    return ""


def _extract_name(text: str) -> str:
    explicit = re.search(r"姓名[:： ]*([^\n]+)", text)
    if explicit:
        value = re.split(r"意向岗位|电话|手机|邮箱|性别|出生|求职意向", explicit.group(1).strip())[0]
        return value.strip(" ：:")
    for line in text.splitlines():
        candidate = line.strip()
        if re.fullmatch(r"[\u4e00-\u9fa5·]{2,8}", candidate):
            return candidate
    return "未知学生"


def _extract_target_job(text: str) -> str:
    match = re.search(r"意向岗位[:： ]*([^\n]+)", text)
    if not match:
        return ""
    value = re.split(r"意向城市|期望薪资|求职类型|比赛经历|项目经历|教育背景|技能|证书|[，,；;]", match.group(1).strip())[0]
    return value.strip(" ：:")


def _extract_school(text: str) -> str:
    """Extract school/university name from resume text."""
    explicit = re.search(r"(学校|院校|大学|School|University)\s*[:：]\s*([^\n|｜，,；;]+)", text, flags=re.IGNORECASE)
    if explicit:
        return explicit.group(2).strip()
    # Look for university name pattern: XX大学/学院
    match = re.search(r"([\u4e00-\u9fff]{2,20}(?:大学|学院|学校))", text)
    if match:
        return match.group(1)
    return ""


def _extract_grade(text: str) -> str:
    """Extract grade level from resume text."""
    match = re.search(r"(年级|Grade)\s*[:：]\s*([^\n|｜，,；;]+)", text, flags=re.IGNORECASE)
    if match:
        return match.group(2).strip()
    # Look for grade patterns like 大三、研二、2024届
    match = re.search(r"(大一|大二|大三|大四|研一|研二|研三|博一|博二|博三|博四|研{0,1}.{0,1}\d{0,1}年级)", text)
    if match:
        return match.group(1)
    return ""


def _extract_graduation_year(text: str) -> str:
    """Extract graduation year from resume text."""
    match = re.search(r"(毕业年份|毕业时间|Graduation)\s*[:：]\s*(\d{4})", text, flags=re.IGNORECASE)
    if match:
        return match.group(2)
    # Look for year patterns like 2025届, 2026年毕业
    match = re.search(r"(\d{4})\s*届", text)
    if match:
        return match.group(1)
    match = re.search(r"(\d{4})\s*年\s*毕业", text)
    if match:
        return match.group(1)
    return ""


def _extract_competitions(text: str) -> list[str]:
    """Extract competitions and honors from resume text."""
    results = []
    # Look for sections with 竞赛/荣誉/奖项/获奖
    section_match = re.search(
        r"(?:竞赛|荣誉|奖项|获奖|比赛|Award|Honor|Competition)[：:]*\s*\n?((?:[-–•*·\d]+[.、)）\s]*.+\n?)+)",
        text, flags=re.IGNORECASE,
    )
    if section_match:
        items = re.findall(r"[-–•*·]\s*(.+)", section_match.group(1))
        results.extend(item.strip() for item in items if item.strip())
    # Also look for inline patterns
    inline = re.findall(r"(?:竞赛|比赛|获奖|荣誉|奖项)[:： ]*(.+?)(?:\n|$)", text)
    for item in inline:
        item = item.strip()
        if item and len(item) < 200:
            results.append(item)
    return results


def _extract_self_evaluation(text: str) -> str:
    """Extract self-evaluation / self-introduction from resume text."""
    # Look for sections like 自我评价/个人总结/自我介绍
    match = re.search(
        r"(?:自我评价|个人总结|自我介绍|个人简介|Self[- ]?Evaluation|Summary)\s*[:：]?\s*\n?((?:[^\n]+\n?)+?)(?=\n\s*\n|\Z|$)",
        text, flags=re.IGNORECASE,
    )
    if match:
        content = match.group(1).strip()
        # Clean up - take first meaningful paragraph
        lines = [line.strip() for line in content.splitlines() if line.strip()]
        if lines:
            return " ".join(lines[:3])  # Take up to 3 lines
    return ""


def _extract_text_from_office_file(file_name: str, content_bytes: bytes) -> str:
    lowered = file_name.lower()
    if lowered.endswith(".pdf"):
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(content_bytes))
            return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
        except Exception as exc:
            logger.warning("Failed to extract PDF text for %s: %s", file_name, exc)
            return ""
    if lowered.endswith(".docx"):
        try:
            from docx import Document

            document = Document(BytesIO(content_bytes))
            lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        lines.append(row_text)
            return "\n".join(lines).strip()
        except Exception as exc:
            logger.warning("Failed to extract DOCX text for %s: %s", file_name, exc)
            return ""
    return ""


def _empty_ocr_result(document_type: str, message: str) -> dict[str, Any]:
    return {
        "raw_text": "",
        "layout_blocks": [],
        "structured_json": {
            "document_type": document_type,
            "name": "未知学生",
            "school": "",
            "major": "",
            "grade": "",
            "graduation_year": "",
            "target_job": "",
            "skills": [],
            "certificates": [],
            "projects": [],
            "internships": [],
            "competitions": [],
            "self_evaluation": "",
            "gpa": None,
            "ocr_warning": message,
        },
    }


class MockOCRProvider(BaseOCRProvider):
    provider_name = "mock"
    SKILL_CANDIDATES = [
        "Python",
        "JavaScript",
        "TypeScript",
        "React",
        "Next.js",
        "FastAPI",
        "PostgreSQL",
        "Redis",
        "SQL",
        "MySQL",
        "Java",
        "Go",
        "C语言",
        "数据结构",
        "数据分析",
        "数据清洗",
        "大数据分析",
        "Docker",
        "Linux",
        "Figma",
        "机器学习",
        "深度学习",
        "PyTorch",
        "Excel",
        "ECharts",
        "Spring Boot",
        "数据可视化",
    ]
    CERTIFICATE_CANDIDATES = [
        "英语四级",
        "英语六级",
        "软件设计师",
        "计算机二级",
        "数据分析师证书",
        "产品经理证书",
        "人工智能工程师",
    ]

    async def parse_document(
        self,
        file_name: str,
        content_bytes: bytes,
        document_type: str = "resume",
        raw_text: Optional[str] = None,
    ) -> dict[str, Any]:
        text = raw_text or _extract_text_from_office_file(file_name, content_bytes)
        if not text:
            lowered = file_name.lower()
            if lowered.endswith((".pdf", ".doc", ".docx", ".png", ".jpg", ".jpeg")):
                raise OCRParseError(
                    "未能从文件中提取可用文字，请使用真实 OCR 或上传可复制文本的 PDF/DOCX。",
                )
            text = content_bytes.decode("utf-8", errors="ignore")
        if not text.strip():
            raise OCRParseError("未能从文件中提取可用文字，文件内容为空。")
        text = _normalize_ocr_text(text)
        skills = _extract_keywords(text, self.SKILL_CANDIDATES)
        certificates = _extract_keywords(text, self.CERTIFICATE_CANDIDATES)
        projects = re.findall(r"(项目|Project)[:： ]*(.+)", text)
        internships = re.findall(r"(实习|Internship)[:： ]*(.+)", text)
        structured = {
            "document_type": document_type,
            "name": _extract_name(text),
            "school": _extract_school(text),
            "major": _extract_major(text),
            "grade": _extract_grade(text),
            "graduation_year": _extract_graduation_year(text),
            "target_job": _extract_target_job(text),
            "skills": skills,
            "certificates": certificates,
            "projects": [item[1].strip() for item in projects],
            "internships": [item[1].strip() for item in internships],
            "competitions": _extract_competitions(text),
            "self_evaluation": _extract_self_evaluation(text),
            "gpa": _extract_gpa(text),
        }
        layout_blocks = [
            {"section": "header", "text": structured["name"]},
            {"section": "skills", "text": ", ".join(skills)},
            {"section": "certificates", "text": ", ".join(certificates)},
            {"section": "experience", "text": "; ".join(structured["internships"] + structured["projects"])},
        ]
        return {
            "raw_text": text,
            "layout_blocks": layout_blocks,
            "structured_json": structured,
        }


class PaddleOCRProvider(BaseOCRProvider):
    provider_name = "paddle"
    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff"}
    RETRYABLE_STATUS_CODES = {429, 500, 502, 503, 504}

    def __init__(
        self,
        service_url: str,
        api_key: str = "",
        *,
        timeout_seconds: float = 60.0,
        max_retries: int = 5,
        retry_base_delay_seconds: float = 2.0,
    ) -> None:
        self.service_url = service_url.rstrip("/")
        self.api_key = api_key
        self.timeout_seconds = timeout_seconds
        self.max_retries = max(1, max_retries)
        self.retry_base_delay_seconds = max(0.0, retry_base_delay_seconds)

    @classmethod
    def _file_type(cls, file_name: str) -> int:
        lowered = file_name.lower()
        return 1 if any(lowered.endswith(ext) for ext in cls.IMAGE_EXTENSIONS) else 0

    @staticmethod
    def _response_error_code(response: httpx.Response) -> int | None:
        try:
            payload = response.json()
        except (ValueError, json.JSONDecodeError):
            return None
        if not isinstance(payload, dict):
            return None
        value = payload.get("errorCode")
        try:
            return int(value) if value is not None else None
        except (TypeError, ValueError):
            return None

    def _should_retry_http_error(self, response: httpx.Response) -> bool:
        if response.status_code in self.RETRYABLE_STATUS_CODES:
            return True
        if self._response_error_code(response) == 10010:
            return True
        body_text = response.text
        return "任务提交队列已满" in body_text or "队列已满" in body_text

    def _retry_delay_seconds(self, attempt: int) -> float:
        return self.retry_base_delay_seconds * attempt

    async def _post_with_retry(
        self,
        client: httpx.AsyncClient,
        *,
        file_name: str,
        payload: dict[str, Any],
        headers: dict[str, str],
    ) -> httpx.Response:
        for attempt in range(1, self.max_retries + 1):
            try:
                logger.info(
                    "PaddleOCR request: POST %s file=%s attempt=%s/%s",
                    self.service_url,
                    file_name,
                    attempt,
                    self.max_retries,
                )
                response = await client.post(
                    self.service_url,
                    json=payload,
                    headers=headers,
                )
                response.raise_for_status()
                return response
            except httpx.HTTPStatusError as exc:
                body_text = exc.response.text[:500]
                if attempt < self.max_retries and self._should_retry_http_error(exc.response):
                    delay = self._retry_delay_seconds(attempt)
                    logger.warning(
                        "PaddleOCR transient HTTP error: file=%s status=%d attempt=%s/%s retry_in=%.1fs body=%s",
                        file_name,
                        exc.response.status_code,
                        attempt,
                        self.max_retries,
                        delay,
                        body_text,
                    )
                    await asyncio.sleep(delay)
                    continue
                raise
            except httpx.TimeoutException:
                if attempt < self.max_retries:
                    delay = self._retry_delay_seconds(attempt)
                    logger.warning(
                        "PaddleOCR timeout: file=%s url=%s attempt=%s/%s retry_in=%.1fs",
                        file_name,
                        self.service_url,
                        attempt,
                        self.max_retries,
                        delay,
                    )
                    await asyncio.sleep(delay)
                    continue
                raise
            except httpx.RequestError as exc:
                if attempt < self.max_retries:
                    delay = self._retry_delay_seconds(attempt)
                    logger.warning(
                        "PaddleOCR connection error: file=%s url=%s attempt=%s/%s retry_in=%.1fs err=%s",
                        file_name,
                        self.service_url,
                        attempt,
                        self.max_retries,
                        delay,
                        exc,
                    )
                    await asyncio.sleep(delay)
                    continue
                raise
        raise OCRServiceError("OCR 鏈嶅姟寮傚父锛岃绋嶅悗閲嶈瘯")

    async def _normalize_layout_parsing_result(
        self,
        file_name: str,
        content_bytes: bytes,
        document_type: str,
        result: dict[str, Any],
    ) -> dict[str, Any]:
        if result.get("raw_text") is not None and result.get("structured_json") is not None:
            return result

        payload = result.get("result") if isinstance(result.get("result"), dict) else result
        layout_results = payload.get("layoutParsingResults") if isinstance(payload, dict) else None
        if not isinstance(layout_results, list):
            raise OCRParseError("OCR 服务返回格式异常，未找到 layoutParsingResults。")

        texts: list[str] = []
        layout_blocks: list[dict[str, Any]] = []
        for index, item in enumerate(layout_results, start=1):
            if not isinstance(item, dict):
                continue
            markdown = item.get("markdown") if isinstance(item.get("markdown"), dict) else {}
            text = markdown.get("text") or item.get("text") or item.get("markdownText")
            if text:
                text_value = str(text).strip()
                texts.append(text_value)
                layout_blocks.append({"section": f"page_{index}", "text": text_value})

        raw_text = "\n\n".join(texts).strip()
        if not raw_text:
            raise OCRParseError("OCR 服务未返回可用文字，请确认文件内容清晰且格式受支持。")

        parsed = await MockOCRProvider().parse_document(
            file_name,
            content_bytes,
            document_type=document_type,
            raw_text=raw_text,
        )
        parsed["layout_blocks"] = layout_blocks or parsed.get("layout_blocks", [])
        return parsed

    async def parse_document(
        self,
        file_name: str,
        content_bytes: bytes,
        document_type: str = "resume",
        raw_text: Optional[str] = None,
    ) -> dict[str, Any]:
        if raw_text is not None:
            logger.info("Using provided raw text for OCR normalization: file=%s", file_name)
            return await MockOCRProvider().parse_document(file_name, content_bytes, document_type, raw_text=raw_text)

        headers = {"Content-Type": "application/json"}
        if self.api_key:
            headers["Authorization"] = f"token {self.api_key}"
        payload = {
            "file": base64.b64encode(content_bytes).decode("ascii"),
            "fileType": self._file_type(file_name),
            "useDocOrientationClassify": False,
            "useDocUnwarping": False,
            "useChartRecognition": False,
        }
        try:
            async with httpx.AsyncClient(timeout=self.timeout_seconds) as client:
                response = await self._post_with_retry(
                    client,
                    file_name=file_name,
                    payload=payload,
                    headers=headers,
                )
                response.raise_for_status()
                result = response.json()
                logger.info("PaddleOCR success: file=%s status=%d", file_name, response.status_code)
                return await self._normalize_layout_parsing_result(file_name, content_bytes, document_type, result)
        except httpx.HTTPStatusError as e:
            logger.error(
                "PaddleOCR HTTP error: file=%s status=%d body=%s",
                file_name, e.response.status_code, e.response.text[:500],
            )
            raise OCRServiceError(
                f"OCR 服务异常 (HTTP {e.response.status_code})，请稍后重试",
            ) from e
        except httpx.TimeoutException as e:
            logger.error("PaddleOCR timeout: file=%s url=%s", file_name, self.service_url)
            raise OCRNetworkError("OCR 服务响应超时，请检查网络后重试") from e
        except httpx.RequestError as e:
            logger.error("PaddleOCR connection error: file=%s url=%s err=%s", file_name, self.service_url, e)
            raise OCRNetworkError(f"无法连接到 OCR 服务，请检查网络后重试") from e
